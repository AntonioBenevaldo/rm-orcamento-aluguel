from copy import deepcopy
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


SOURCE = Path(
    r"C:\Users\Benevaldo\Downloads"
    r"\Relatorio_Revisado_Orcamento_Imobiliario_RM (2).docx"
)
OUTPUT = Path("docs/Relatorio_Final_Orcamento_Imobiliario_RM.docx")


def set_paragraph_text(paragraph, text):
    """Replace visible text while retaining the paragraph and first-run formatting."""
    saved_rpr = None
    if paragraph.runs and paragraph.runs[0]._r.rPr is not None:
        saved_rpr = deepcopy(paragraph.runs[0]._r.rPr)
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)
    run = paragraph.add_run(text)
    if saved_rpr is not None:
        run._r.insert(0, saved_rpr)


def insert_after(paragraph, text, style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if style:
        new_paragraph.style = style
    new_paragraph.add_run(text)
    return new_paragraph


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def copy_row_format(source_row, target_row):
    if source_row._tr.trPr is not None:
        target_row._tr.insert(0, deepcopy(source_row._tr.trPr))
    for source_cell, target_cell in zip(source_row.cells, target_row.cells):
        if source_cell._tc.tcPr is not None:
            target_cell._tc.insert(0, deepcopy(source_cell._tc.tcPr))
        for source_p, target_p in zip(source_cell.paragraphs, target_cell.paragraphs):
            if source_p._p.pPr is not None:
                target_p._p.insert(0, deepcopy(source_p._p.pPr))


def set_cell(cell, text):
    p = cell.paragraphs[0]
    set_paragraph_text(p, text)


def add_table_row_like(table, values, template_index=-1):
    template = table.rows[template_index]
    row = table.add_row()
    copy_row_format(template, row)
    for cell, value in zip(row.cells, values):
        set_cell(cell, value)
    return row


def update_document():
    doc = Document(SOURCE)
    p = doc.paragraphs

    replacements = {
        97: (
            "A solução utiliza uma arquitetura monolítica modular em camadas. "
            "O mesmo núcleo de domínio, serviços, persistência e exportação é "
            "acessado por duas interfaces: a aplicação web em Streamlit e a "
            "aplicação interativa de terminal. Essa separação facilita testes, "
            "manutenção e evolução sem duplicar as regras de negócio."
        ),
        100: (
            "Streamlit foi escolhido para fornecer a interface web no navegador, "
            "enquanto main.py oferece uma segunda forma de uso diretamente no terminal."
        ),
        104: (
            "As dependências seguem a direção apresentação (Streamlit ou terminal), "
            "aplicação, domínio e persistência."
        ),
        106: (
            "A interface escolhida - navegador ou terminal - recebe os dados e "
            "realiza validações básicas."
        ),
        108: (
            "O domínio produz itens e totais sem depender do Streamlit nem do terminal."
        ),
        142: (
            "O projeto foi estruturado como um monólito modular em Python com duas "
            "interfaces funcionais: Streamlit, acessível no navegador em "
            "http://localhost:8501, e terminal, executada por main.py. Ambas delegam "
            "os cálculos aos mesmos serviços e modelos orientados a objetos, enquanto "
            "a persistência registra os resultados em SQLite."
        ),
        152: (
            "Validar regras, banco de dados, interfaces web e terminal e exportação "
            "mediante testes automatizados."
        ),
        216: "app.py - inicialização da aplicação web em Streamlit",
        217: "main.py - aplicação interativa executada diretamente no terminal",
        218: "pages/ - páginas da interface Streamlit",
        219: "models/ - entidades e objetos de domínio compartilhados pelas duas interfaces",
        220: "services/ - cálculo e exportação compartilhados pelas duas interfaces",
        221: "database/ - conexão, consultas e transações SQLite",
        222: "tests/ - 18 testes unitários, de integração e das interfaces",
        228: (
            "Um único núcleo de serviços atende às interfaces Streamlit e terminal, "
            "eliminando duplicidade de regras e facilitando manutenção e testes."
        ),
        230: "O usuário informa os dados pela interface Streamlit ou pela aplicação de terminal.",
        231: "A interface valida os campos e envia os dados ao CalculoService.",
        262: (
            "Foram executados 18 testes automatizados envolvendo regras de negócio, "
            "persistência, integridade, interfaces web e terminal e exportação. "
            "Toda a suíte foi aprovada, inclusive os três cenários oficiais."
        ),
        271: (
            "O CSV deve conter todos os 12 meses e refletir os mesmos valores exibidos "
            "na interface web ou no terminal."
        ),
        273: "17. Execução pelas interfaces web e terminal",
        274: (
            "No Windows, abra o PowerShell ou o terminal do VS Code na pasta do projeto. "
            "Prepare o ambiente e execute os testes:"
        ),
        280: "python main.py",
        281: (
            "A execução acima abre a versão interativa no próprio terminal. Para usar "
            "a versão web, execute o comando abaixo; o navegador poderá ser aberto em "
            "http://localhost:8501."
        ),
        283: "app.py, main.py e demais arquivos .py do projeto.",
        286: (
            "README.md com objetivo, instalação, execução das duas interfaces, testes "
            "e link do vídeo."
        ),
        296: (
            "Explicar brevemente as tecnologias, a estrutura orientada a objetos e as "
            "duas formas de execução."
        ),
        297: (
            "Demonstrar um orçamento na interface Streamlit e repetir a operação no "
            "terminal para comprovar que ambas funcionam."
        ),
        304: (
            "A utilização de centavos inteiros elimina inconsistências monetárias entre "
            "o código, o diagrama de classes e o banco SQLite. Os 18 testes automatizados "
            "aprovados fornecem evidências de que os cenários oficiais, a integridade do "
            "banco, a exportação e as duas interfaces funcionam conforme especificado."
        ),
    }
    for index, text in replacements.items():
        set_paragraph_text(p[index], text)

    # Insert the Streamlit command immediately after the terminal explanation.
    inserted = insert_after(p[281], "python -m streamlit run app.py", "Intense Quote")

    # Add convenient Windows launchers after the Streamlit command.
    note = insert_after(
        inserted,
        "Atalhos equivalentes no Windows: executar_terminal.bat para o terminal e "
        "executar_aplicacao.bat para o Streamlit.",
        "Normal",
    )
    insert_after(
        note,
        "As duas versões utilizam as mesmas classes, serviços, banco SQLite e rotina "
        "de exportação CSV; portanto, qualquer correção de regra beneficia as duas.",
        "Normal",
    )

    # Functional requirement: the assignment can be demonstrated in both modes.
    add_table_row_like(
        doc.tables[5],
        (
            "RF09",
            "Disponibilizar o orçamento por duas formas funcionais: interface web "
            "Streamlit e aplicação interativa de terminal.",
        ),
    )

    # Architecture table.
    set_cell(doc.tables[7].cell(1, 1), "app.py, pages/ e main.py")
    set_cell(
        doc.tables[7].cell(1, 2),
        "Coletar entradas e apresentar resultados no navegador ou no terminal.",
    )

    # Technology table.
    set_cell(
        doc.tables[13].cell(1, 1),
        "Streamlit >= 1.48 e terminal Python",
    )
    set_cell(
        doc.tables[13].cell(1, 2),
        "Duas interfaces para o mesmo núcleo da aplicação.",
    )

    # Test evidence table.
    add_table_row_like(
        doc.tables[16],
        (
            "Interface terminal: fluxo completo com entradas simuladas",
            "Orçamento calculado e apresentado sem erro",
            "Orçamento calculado e apresentado sem erro",
            "Aprovado",
        ),
    )
    add_table_row_like(
        doc.tables[16],
        (
            "Suíte automatizada completa",
            "18 testes aprovados",
            "18 testes aprovados",
            "Aprovado",
        ),
    )

    # Deliverables table.
    set_cell(
        doc.tables[17].cell(2, 1),
        "Pasta compactada com código funcional nas versões Streamlit e terminal, "
        "além do link do GitHub.",
    )

    # Metadata and update-field hint for Word.
    doc.core_properties.title = (
        "Relatório Final - Sistema de Orçamento Imobiliário R.M"
    )
    doc.core_properties.subject = (
        "Aplicações Streamlit e terminal com modelo físico SQLite"
    )
    doc.core_properties.modified = datetime.now()
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT.resolve())


if __name__ == "__main__":
    update_document()
